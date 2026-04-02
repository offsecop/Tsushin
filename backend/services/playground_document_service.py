"""
Phase 14.2: Playground Document Service
Handles document uploads, processing, and knowledge base management for Playground conversations.
"""

import os
import logging
import uuid
import json
from typing import Dict, List, Optional, Any
from datetime import datetime
from pathlib import Path

from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)


class PlaygroundDocumentService:
    """
    Service for handling document attachments in Playground.

    Key responsibilities:
    - Store uploaded documents
    - Extract text from various formats (PDF, TXT, CSV, JSON, XLSX, DOCX, MD, RTF)
    - Chunk and embed documents
    - Store in tenant-isolated ChromaDB collections
    - Clean up documents when conversation is cleared
    """

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.txt': 'txt',
        '.csv': 'csv',
        '.json': 'json',
        '.xlsx': 'xlsx',
        '.xls': 'xlsx',
        '.docx': 'docx',
        '.doc': 'docx',
        '.md': 'md',
        '.markdown': 'md',
        '.rtf': 'rtf',
    }

    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

    def __init__(self, db: Session):
        self.db = db
        self.logger = logging.getLogger(__name__)

    def get_storage_path(self, tenant_id: str, user_id: int) -> str:
        """Get storage path for user's playground documents."""
        import settings
        base_path = getattr(settings, 'DATA_DIR', 'data')
        path = os.path.join(base_path, 'playground_docs', tenant_id, str(user_id))
        os.makedirs(path, exist_ok=True)
        return path

    def get_collection_name(self, tenant_id: str, user_id: int, agent_id: int) -> str:
        """Get ChromaDB collection name for user's conversation."""
        return f"playground_{tenant_id}_{user_id}_{agent_id}"

    async def upload_document(
        self,
        tenant_id: str,
        user_id: int,
        agent_id: int,
        file_data: bytes,
        filename: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        embedding_model: str = "all-MiniLM-L6-v2"
    ) -> Dict[str, Any]:
        """
        Upload and process a document.

        Args:
            tenant_id: Tenant ID for isolation
            user_id: User ID
            agent_id: Agent ID for the conversation
            file_data: File bytes
            filename: Original filename
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
            embedding_model: Model for embeddings

        Returns:
            Dict with document info and status
        """
        from models import PlaygroundDocument, PlaygroundDocumentChunk

        try:
            # Validate file extension
            ext = Path(filename).suffix.lower()
            if ext not in self.SUPPORTED_EXTENSIONS:
                return {
                    "status": "error",
                    "error": f"Unsupported file type: {ext}. Supported: {', '.join(self.SUPPORTED_EXTENSIONS.keys())}"
                }

            # Validate file size
            if len(file_data) > self.MAX_FILE_SIZE:
                return {
                    "status": "error",
                    "error": f"File too large. Maximum size is {self.MAX_FILE_SIZE // (1024*1024)} MB"
                }

            # Generate conversation ID
            conversation_id = f"playground_{user_id}_{agent_id}"

            # Save file
            storage_path = self.get_storage_path(tenant_id, user_id)
            doc_id = str(uuid.uuid4())
            file_path = os.path.join(storage_path, f"{doc_id}{ext}")

            with open(file_path, 'wb') as f:
                f.write(file_data)

            self.logger.info(f"Saved document to {file_path}")

            # Create document record
            doc = PlaygroundDocument(
                tenant_id=tenant_id,
                user_id=user_id,
                agent_id=agent_id,
                conversation_id=conversation_id,
                document_name=filename,
                document_type=self.SUPPORTED_EXTENSIONS[ext],
                file_path=file_path,
                file_size_bytes=len(file_data),
                embedding_model=embedding_model,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                status="processing"
            )
            self.db.add(doc)
            self.db.commit()
            self.db.refresh(doc)

            # Process document asynchronously
            try:
                await self._process_document(doc, chunk_size, chunk_overlap, embedding_model)
                doc.status = "completed"
                doc.processed_date = datetime.utcnow()
            except Exception as e:
                self.logger.error(f"Document processing failed: {e}", exc_info=True)
                doc.status = "failed"
                doc.error_message = str(e)

            self.db.commit()

            return {
                "status": "success",
                "document": {
                    "id": doc.id,
                    "name": doc.document_name,
                    "type": doc.document_type,
                    "size_bytes": doc.file_size_bytes,
                    "num_chunks": doc.num_chunks,
                    "status": doc.status,
                    "error": doc.error_message
                }
            }

        except Exception as e:
            self.logger.error(f"Document upload failed: {e}", exc_info=True)
            return {
                "status": "error",
                "error": str(e)
            }

    async def _process_document(
        self,
        doc,
        chunk_size: int,
        chunk_overlap: int,
        embedding_model: str
    ):
        """Process document: extract text, chunk, and embed."""
        from models import PlaygroundDocumentChunk

        # Extract text based on document type
        text = await self._extract_text(doc.file_path, doc.document_type)

        if not text:
            raise ValueError("No text content extracted from document")

        # Chunk the text
        chunks = self._chunk_text(text, chunk_size, chunk_overlap)

        if not chunks:
            raise ValueError("No chunks created from document")

        self.logger.info(f"Created {len(chunks)} chunks from document")

        # Store chunks in database
        for i, chunk_text in enumerate(chunks):
            chunk = PlaygroundDocumentChunk(
                document_id=doc.id,
                chunk_index=i,
                content=chunk_text,
                char_count=len(chunk_text),
                metadata_json={
                    "document_name": doc.document_name,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
            )
            self.db.add(chunk)

        doc.num_chunks = len(chunks)
        self.db.commit()

        # Store embeddings in ChromaDB
        await self._store_embeddings(doc, chunks, embedding_model)

    async def _extract_text(self, file_path: str, doc_type: str) -> str:
        """Extract text from document based on type."""
        try:
            if doc_type == 'txt' or doc_type == 'md':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()

            elif doc_type == 'pdf':
                try:
                    import pypdf
                    with open(file_path, 'rb') as f:
                        reader = pypdf.PdfReader(f)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                        return text
                except ImportError:
                    self.logger.warning("pypdf not installed, trying pdfplumber")
                    try:
                        import pdfplumber
                        with pdfplumber.open(file_path) as pdf:
                            text = ""
                            for page in pdf.pages:
                                text += (page.extract_text() or "") + "\n"
                            return text
                    except ImportError:
                        raise ImportError("Install pypdf or pdfplumber for PDF support")

            elif doc_type == 'docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    return text
                except ImportError:
                    raise ImportError("Install python-docx for DOCX support")

            elif doc_type == 'csv':
                import csv
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    reader = csv.reader(f)
                    rows = [",".join(row) for row in reader]
                    return "\n".join(rows)

            elif doc_type == 'json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return json.dumps(data, indent=2, ensure_ascii=False)

            elif doc_type == 'xlsx':
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(file_path, read_only=True)
                    text = ""
                    for sheet in wb.worksheets:
                        for row in sheet.iter_rows(values_only=True):
                            row_text = "\t".join([str(cell) if cell else "" for cell in row])
                            text += row_text + "\n"
                    return text
                except ImportError:
                    raise ImportError("Install openpyxl for XLSX support")

            elif doc_type == 'rtf':
                try:
                    from striprtf.striprtf import rtf_to_text
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        rtf_content = f.read()
                        return rtf_to_text(rtf_content)
                except ImportError:
                    raise ImportError("Install striprtf for RTF support")

            else:
                # Try to read as plain text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()

        except Exception as e:
            self.logger.error(f"Text extraction failed for {doc_type}: {e}")
            raise

    def _chunk_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """Split text into overlapping chunks."""
        if not text:
            return []

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = start + chunk_size
            chunk = text[start:end]

            # Try to break at sentence or word boundary
            if end < text_len:
                # Look for sentence boundary
                for sep in ['. ', '.\n', '! ', '!\n', '? ', '?\n', '\n\n']:
                    last_sep = chunk.rfind(sep)
                    if last_sep > chunk_size * 0.5:  # Only if we're past halfway
                        chunk = chunk[:last_sep + len(sep)]
                        break

            chunks.append(chunk.strip())
            start = start + len(chunk) - overlap

            # Prevent infinite loop
            if start <= 0 and len(chunks) > 1:
                break

        return [c for c in chunks if c]  # Remove empty chunks

    async def _store_embeddings(self, doc, chunks: List[str], embedding_model: str):
        """
        Store chunk embeddings in ChromaDB.

        BUG-001 Fix: Uses shared embedding service with batched processing
        to prevent OOM crashes on large documents.
        """
        try:
            import chromadb
            from agent.memory.embedding_service import get_shared_embedding_service
            import settings

            # Get ChromaDB client
            persist_dir = getattr(settings, 'CHROMA_PERSIST_DIR', 'data/chroma')
            client = chromadb.PersistentClient(path=persist_dir)

            # Get or create collection for this conversation
            collection_name = self.get_collection_name(doc.tenant_id, doc.user_id, doc.agent_id)
            collection = client.get_or_create_collection(
                name=collection_name,
                metadata={"hnsw:space": "cosine"}
            )

            # BUG-001 Fix: Use shared service with batched processing (async)
            embedding_service = get_shared_embedding_service(embedding_model)
            embeddings = await embedding_service.embed_batch_chunked_async(chunks, batch_size=50)

            # Validate we got embeddings for all chunks
            if len(embeddings) != len(chunks):
                self.logger.warning(
                    f"Embedding count mismatch: {len(embeddings)} embeddings for {len(chunks)} chunks"
                )
                # Only process chunks we have embeddings for
                chunks = chunks[:len(embeddings)]

            # Add to collection
            ids = [f"{doc.id}_{i}" for i in range(len(chunks))]
            metadatas = [
                {
                    "document_id": doc.id,
                    "document_name": doc.document_name,
                    "chunk_index": i,
                    "conversation_id": doc.conversation_id
                }
                for i in range(len(chunks))
            ]

            collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=chunks,
                metadatas=metadatas
            )

            self.logger.info(f"Stored {len(chunks)} embeddings in collection {collection_name}")

        except Exception as e:
            self.logger.error(f"Failed to store embeddings: {e}", exc_info=True)
            # Don't fail the whole upload if embedding storage fails

    async def get_documents(
        self,
        tenant_id: str,
        user_id: int,
        agent_id: int
    ) -> List[Dict[str, Any]]:
        """Get all documents for a conversation."""
        from models import PlaygroundDocument

        conversation_id = f"playground_{user_id}_{agent_id}"

        docs = self.db.query(PlaygroundDocument).filter(
            PlaygroundDocument.tenant_id == tenant_id,
            PlaygroundDocument.user_id == user_id,
            PlaygroundDocument.agent_id == agent_id,
            PlaygroundDocument.conversation_id == conversation_id
        ).order_by(PlaygroundDocument.upload_date.desc()).all()

        return [
            {
                "id": doc.id,
                "name": doc.document_name,
                "type": doc.document_type,
                "size_bytes": doc.file_size_bytes,
                "num_chunks": doc.num_chunks,
                "status": doc.status,
                "error": doc.error_message,
                "upload_date": doc.upload_date.isoformat() if doc.upload_date else None
            }
            for doc in docs
        ]

    async def delete_document(
        self,
        tenant_id: str,
        user_id: int,
        doc_id: int
    ) -> Dict[str, Any]:
        """Delete a document and its embeddings."""
        from models import PlaygroundDocument, PlaygroundDocumentChunk

        try:
            doc = self.db.query(PlaygroundDocument).filter(
                PlaygroundDocument.id == doc_id,
                PlaygroundDocument.tenant_id == tenant_id,
                PlaygroundDocument.user_id == user_id
            ).first()

            if not doc:
                return {"status": "error", "error": "Document not found"}

            # Delete chunks from database
            self.db.query(PlaygroundDocumentChunk).filter(
                PlaygroundDocumentChunk.document_id == doc_id
            ).delete()

            # Delete embeddings from ChromaDB
            try:
                await self._delete_embeddings(doc)
            except Exception as e:
                self.logger.warning(f"Failed to delete embeddings: {e}")

            # Delete file
            try:
                if os.path.exists(doc.file_path):
                    os.remove(doc.file_path)
            except Exception as e:
                self.logger.warning(f"Failed to delete file: {e}")

            # Delete document record
            self.db.delete(doc)
            self.db.commit()

            return {"status": "success", "message": "Document deleted"}

        except Exception as e:
            self.logger.error(f"Failed to delete document: {e}", exc_info=True)
            return {"status": "error", "error": str(e)}

    async def _delete_embeddings(self, doc):
        """Delete embeddings from ChromaDB."""
        try:
            import chromadb
            import settings

            persist_dir = getattr(settings, 'CHROMA_PERSIST_DIR', 'data/chroma')
            client = chromadb.PersistentClient(path=persist_dir)

            collection_name = self.get_collection_name(doc.tenant_id, doc.user_id, doc.agent_id)

            try:
                collection = client.get_collection(name=collection_name)
                # Delete by document ID prefix
                ids_to_delete = [f"{doc.id}_{i}" for i in range(doc.num_chunks)]
                if ids_to_delete:
                    collection.delete(ids=ids_to_delete)
            except Exception:
                pass  # Collection might not exist

        except Exception as e:
            self.logger.warning(f"Failed to delete embeddings: {e}")

    async def clear_all_documents(
        self,
        tenant_id: str,
        user_id: int,
        agent_id: int
    ) -> Dict[str, Any]:
        """Clear all documents for a conversation."""
        from models import PlaygroundDocument

        try:
            conversation_id = f"playground_{user_id}_{agent_id}"

            docs = self.db.query(PlaygroundDocument).filter(
                PlaygroundDocument.tenant_id == tenant_id,
                PlaygroundDocument.user_id == user_id,
                PlaygroundDocument.agent_id == agent_id,
                PlaygroundDocument.conversation_id == conversation_id
            ).all()

            deleted_count = 0
            for doc in docs:
                result = await self.delete_document(tenant_id, user_id, doc.id)
                if result.get("status") == "success":
                    deleted_count += 1

            return {
                "status": "success",
                "message": f"Deleted {deleted_count} documents"
            }

        except Exception as e:
            self.logger.error(f"Failed to clear documents: {e}", exc_info=True)
            return {"status": "error", "error": str(e)}

    async def search_documents(
        self,
        tenant_id: str,
        user_id: int,
        agent_id: int,
        query: str,
        max_results: int = 5
    ) -> List[Dict[str, Any]]:
        """Search documents using semantic search."""
        try:
            import chromadb
            from agent.memory.embedding_service import get_shared_embedding_service
            import settings

            persist_dir = getattr(settings, 'CHROMA_PERSIST_DIR', 'data/chroma')
            client = chromadb.PersistentClient(path=persist_dir)

            collection_name = self.get_collection_name(tenant_id, user_id, agent_id)

            try:
                collection = client.get_collection(name=collection_name)
            except Exception:
                return []  # No documents uploaded yet

            # Generate query embedding
            embedding_service = get_shared_embedding_service("all-MiniLM-L6-v2")
            query_embedding = await embedding_service.embed_text_async(query)

            # Search
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=max_results
            )

            if not results or not results.get('documents'):
                return []

            documents = results.get('documents', [[]])[0]
            metadatas = results.get('metadatas', [[]])[0]
            distances = results.get('distances', [[]])[0]

            return [
                {
                    "content": doc,
                    "metadata": meta,
                    "similarity": 1 - dist  # Convert distance to similarity
                }
                for doc, meta, dist in zip(documents, metadatas, distances)
            ]

        except Exception as e:
            self.logger.error(f"Search failed: {e}", exc_info=True)
            return []
